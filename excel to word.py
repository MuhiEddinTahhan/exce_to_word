import pandas as pd
from docx import Document

def excel_to_word(excel_file, sheet_name, row_indices, output_file):
    # Read the Excel file
    df = pd.read_excel(excel_file, sheet_name=sheet_name)

    # Create a Word document
    doc = Document()
    doc.add_heading(f"Data from {excel_file}", level=1)

    # Add specific rows to the Word document
    for row_index in row_indices:
        row = df.iloc[row_index]
        # Add a new paragraph for each row
        paragraph = doc.add_paragraph()

        # Add each cell value to the corresponding column in the Word document
        for column_name, cell_value in row.items():
            run = paragraph.add_run(f"{column_name}: ")
            run.bold = True  # Bold column name
            paragraph.add_run(str(cell_value))  # Regular cell value
            paragraph.add_run("\n")  # Add a new line after each cell value

    # Save the Word document
    doc.save(output_file)
    print(f"Data from '{excel_file}' has been saved to '{output_file}'.")

if __name__ == "__main__":
    # Replace these file paths with your actual paths
    excel_file_path = r"your/path/here.xlsx"
    sheet_name = "Sheet1"  # Replace with the sheet name containing the data
    row_indices = [0, 2, 4]  # Replace with the row indices you want to include
    output_file_path = r"your/path/here.docx"

    excel_to_word(excel_file_path, sheet_name, row_indices, output_file_path)
